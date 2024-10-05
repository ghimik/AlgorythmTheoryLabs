import docx
import openpyxl

from decorators import validate_positive_numbers, log_method_call
from lab1.particles import Particle
from particles import CustomParticle
from particles import Electron
from particles import Proton
from particles import Neutron


def _represent_particle_(particle: Particle):
    return {
            "Specific Charge (C/kg)": particle.specific_charge,
            "Compton Wavelength (m)": particle.compton_wavelength,
        }


class ParticleCalculator:
    """Класс для вычисления свойств частиц."""

    def __init__(self):
        self.particles: dict[str: Particle] = {
            "electron": Electron(),
            "proton": Proton(),
            "neutron": Neutron()
        }

    @log_method_call
    def display_results(self):
        """Отображение результатов расчета для всех частиц."""
        results = {}
        for name, particle in self.particles.items():
            results[name] = _represent_particle_(particle)
        return results

    @log_method_call
    @validate_positive_numbers
    def add_custom_particle(self, mass, charge):
        """Добавление кастомной частицы в расчеты."""
        custom_particle = CustomParticle(mass, charge)
        self.particles["custom"] = custom_particle

    @log_method_call
    def save_to_docx(self, data):
        """Сохранение результатов в формате .docx."""
        doc = docx.Document()
        doc.add_heading('Results of Particle Calculations', 0)

        for particle, results in data.items():
            doc.add_paragraph(f"{particle.capitalize()}:")
            for key, value in results.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph()
        doc.save('results.docx')

    @log_method_call
    def save_to_xlsx(self, data):
        """Сохранение результатов в формате .xlsx."""
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Particle Data"

        sheet.append(["Particle", "Property", "Value"])

        for particle, results in data.items():
            for key, value in results.items():
                sheet.append([particle.capitalize(), key, value])
        workbook.save('results.xlsx')

    @log_method_call
    def run(self):
        """Запуск программы."""
        print("Welcome to Particle Calculator")
        print("Calculating specific charge and Compton wavelength for particles.\n")

        results = self.display_results()

        for name, result in results.items():
            print(f"{name.capitalize()}:")
            for key, value in result.items():
                print(f"  {key}: {value}")
            print()

        custom_choice = input("Would you like to create a custom particle? (yes/no): ").strip().lower()
        if custom_choice == "yes":
            mass = float(input("Enter mass (kg): "))
            charge = float(input("Enter charge (C): "))
            self.add_custom_particle(mass, charge)
            particle: Particle = self.particles['custom']
            print(f"Custom particle created: {particle}")
            for k,v in _represent_particle_(particle).items():
                print(f"{k}: {v}")

        save_choice = input("Would you like to save results? (docx/xlsx/no): ").strip().lower()
        if save_choice == "docx":
            self.save_to_docx(results)
            print("Results saved to results.docx")
        elif save_choice == "xlsx":
            self.save_to_xlsx(results)
            print("Results saved to results.xlsx")
        elif save_choice == "no":
            print("Results not saved.")
        else:
            print("Unknown command. Only 'docx', 'xlsx', 'no'. ")



